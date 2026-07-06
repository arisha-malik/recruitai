import os
import pdfplumber
from docx import Document

class TextExtractionError(Exception):
    """Base exception for resume text extraction failures."""
    pass

class UnsupportedFileTypeError(TextExtractionError):
    """Raised when the uploaded file type is not supported."""
    pass

class EmptyFileError(TextExtractionError):
    """Raised when the extracted text is empty or only whitespace."""
    pass

class FileCorruptedError(TextExtractionError):
    """Raised when the file cannot be opened/parsed."""
    pass

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file using pdfplumber."""
    text_content = []
    try:
        with pdfplumber.open(file_path) as pdf:
            if not pdf.pages:
                raise EmptyFileError("PDF file contains no pages.")
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
    except EmptyFileError:
        raise
    except Exception as e:
        raise FileCorruptedError(f"Failed to parse PDF file. It might be corrupted: {str(e)}")
        
    full_text = "\n".join(text_content).strip()
    if not full_text:
        raise EmptyFileError("No readable text could be extracted from this PDF.")
    return full_text

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Also extract table text if present
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        paragraphs.append(cell_text)
                        
    except Exception as e:
        raise FileCorruptedError(f"Failed to parse DOCX file. It might be corrupted: {str(e)}")
        
    full_text = "\n".join(paragraphs).strip()
    if not full_text:
        raise EmptyFileError("No readable text could be extracted from this DOCX.")
    return full_text

def extract_text_from_file(file_path: str, original_filename: str) -> str:
    """
    Dispatcher to extract plain text from supported document types based on extension.
    """
    if not os.path.exists(file_path):
        raise TextExtractionError(f"File not found at path: {file_path}")
        
    ext = os.path.splitext(original_filename.lower())[1]
    
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        # Word documents (.docx)
        # Note: python-docx only supports openXML (.docx). For old binary .doc, we treat it similarly or let python-docx try (or throw unsupported).
        # We raise UnsupportedFileTypeError if it is a binary .doc and python-docx fails.
        try:
            return extract_text_from_docx(file_path)
        except FileCorruptedError:
            if ext == ".doc":
                raise UnsupportedFileTypeError(
                    "Legacy binary .doc files are not supported. Please convert to .docx or .pdf."
                )
            raise
    else:
        raise UnsupportedFileTypeError(
            f"Unsupported file type '{ext}'. Only PDF and DOCX files are supported."
        )
