import os
import sys
import uuid
import json
import shutil
from unittest.mock import patch, MagicMock

# Add backend directory to path
backend_path = r"c:\Users\HP\Documents\VS Code\AATA\backend"
sys.path.insert(0, backend_path)

# Set env variables before importing app
os.environ["DATABASE_URL"] = "sqlite:///verify_mvp_no_qdrant.db"
os.environ["QDRANT_ENABLED"] = "false"
os.environ["RUN_JOBS_SYNC"] = "true"

from fastapi.testclient import TestClient
from app.main import app
from app.database import Base, engine, SessionLocal
from app.models.user import User, UserRole
from app.models.candidate import Candidate, Resume, ParsingStatus, ParsedResumeData
from app.models.job import Job, MatchingStatus, JobEmbeddingStatus
from app.models.application import MatchResult, Application, ApplicationStatus
from app.models.audit import RecruitmentEvent

# Helper to create a physical DOCX file for testing text extraction
def create_test_docx(filename: str):
    from docx import Document
    doc = Document()
    doc.add_paragraph("Rahul Sharma")
    doc.add_paragraph("Email: rahul.sharma@gmail.com")
    doc.add_paragraph("Phone: +919876543210")
    doc.add_paragraph("Skills: Python, FastAPI, PostgreSQL")
    doc.save(filename)

# Re-create database schema on our temporary SQLite file
Base.metadata.drop_all(bind=engine)
Base.metadata.create_all(bind=engine)

client = TestClient(app)

# Mocks
MOCK_RESUME_PARSE_DATA = {
    "full_name": "Rahul Sharma",
    "email": "rahul.sharma@gmail.com",
    "mobile_number": "+919876543210",
    "total_experience_years": 4.5,
    "technical_skills": ["Python", "FastAPI", "PostgreSQL"],
    "education": [
        {
            "institution": "IIT",
            "degree": "B.Tech",
            "field_of_study": "Computer Science",
            "start_year": "2018",
            "end_year": "2022"
        }
    ],
    "certifications": ["AWS Solutions Architect"],
    "current_company": "Tech Corp",
    "current_location": "Remote, US",
    "notice_period": "Immediate",
    "work_experience": [
        {
            "company": "Tech Corp",
            "role": "Software Engineer",
            "start_date": "2022-06",
            "end_date": "Present",
            "description": "Building FastAPI apps"
        }
    ],
    "projects": [
        {
            "title": "RecruitAI Platform",
            "description": "Resume parsing backend",
            "technologies": ["FastAPI", "PostgreSQL"]
        }
    ]
}

MOCK_MATCH_EVALUATION = {
    "candidate_id": "Rahul-Sharma-UUID-Placeholder",
    "match_percentage": 92.0,
    "matched_skills": ["Python", "FastAPI"],
    "missing_skills": ["AWS"],
    "skill_match_analysis": "Candidate matches core programming requirements.",
    "experience_analysis": "Has 4.5 years experience which aligns with Mid-Senior requirements.",
    "experience_delta": "Meets experience requirement",
    "location_fit": "Matches Remote US location requirement.",
    "notice_period_fit": "Immediate availability fits notice requirements.",
    "strengths": ["FastAPI experience", "AWS certification"],
    "concerns": ["None"],
    "final_recommendation": "SHORTLIST",
    "summary": "Solid candidate."
}

@patch("app.services.llm_service.LLMService._call_gemini")
@patch("app.services.matching_service.MatchingLLMService._call_gemini")
def run_tests(mock_match_gemini, mock_parse_gemini):
    print("--- STARTING REVISED MVP (NO QDRANT) INTEGRATION TESTS ---")
    
    # Configure mock returns
    mock_parse_gemini.return_value = MOCK_RESUME_PARSE_DATA
    mock_match_gemini.return_value = json.dumps(MOCK_MATCH_EVALUATION)
    
    db = SessionLocal()

    # 1. Sign up admin recruiter
    print("\n1. Creating Recruiter user...")
    r = client.post("/api/v1/auth/signup", json={
        "email": "recruiter@recruitai.com",
        "password": "recruiterpass",
        "first_name": "Jane",
        "last_name": "Doe",
        "role": "RECRUITER"
    })
    assert r.status_code == 201
    
    # Login to obtain JWT
    r = client.post("/api/v1/auth/login", data={"username": "recruiter@recruitai.com", "password": "recruiterpass"})
    assert r.status_code == 200
    token = r.json()["access_token"]
    headers = {"Authorization": f"Bearer {token}"}
    print("Recruiter logged in successfully.")

    # 2. Upload and Parse Candidate Resume without Qdrant
    print("\n2. Testing Resume Upload & Parse without Qdrant...")
    # Trigger presigned url (creates candidate & resume record)
    r = client.post("/api/v1/uploads/presigned-url", json={
        "first_name": "Rahul",
        "last_name": "Sharma",
        "email": "rahul.sharma@gmail.com",
        "file_name": "resume.docx",
        "file_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    }, headers=headers)
    assert r.status_code == 200
    presign_res = r.json()
    cand_id = presign_res["candidate_id"]
    resume_id = presign_res["resume_id"]
    
    # Let's write a mock docx file in uploads folder
    os.makedirs(f"uploads/resumes/{resume_id}", exist_ok=True)
    mock_file_path = f"uploads/resumes/{resume_id}/resume.docx"
    create_test_docx(mock_file_path)
        
    # Trigger parsing
    print("Triggering parse...")
    r = client.post(f"/api/v1/resumes/{resume_id}/parse", headers=headers)
    assert r.status_code == 202
    
    # Retrieve resume status (should be PARSED)
    r = client.get(f"/api/v1/resumes/{resume_id}/status", headers=headers)
    assert r.status_code == 200
    status_res = r.json()
    assert status_res["status"] == "PARSED"
    print("Success: Resume parsing completed successfully with Qdrant disabled.")
    
    # Verify candidate profile details are updated from Gemini parse response
    candidate = db.query(Candidate).filter(Candidate.id == cand_id).first()
    assert candidate.first_name == "Rahul"
    assert candidate.last_name == "Sharma"
    assert candidate.total_experience_years == 4.5
    assert "FastAPI" in candidate.skills
    print(f"Success: Candidate profile successfully updated with skills: {candidate.skills}")
    
    # Check that events RESUME_PARSING_STARTED and RESUME_PARSED are logged
    # and CANDIDATE_VECTOR_CREATED was NOT logged (since Qdrant is disabled)
    events = db.query(RecruitmentEvent).filter(RecruitmentEvent.candidate_id == cand_id).all()
    event_types = [e.event_type for e in events]
    print(f"Recorded Events: {event_types}")
    assert "RESUME_PARSING_STARTED" in event_types
    assert "RESUME_PARSED" in event_types
    assert "CANDIDATE_VECTOR_CREATED" not in event_types
    print("Success: Events checked. Vector creation event correctly skipped.")

    # 3. Basic Candidate Matching without Qdrant
    print("\n3. Testing Job Matching without Qdrant...")
    # Create a Job
    r = client.post("/api/v1/jobs/", json={
        "title": "Senior Python Backend Engineer (FastAPI/Python)",
        "department": "Engineering",
        "location": "Remote, US",
        "experience_level": "Senior",
        "employment_type": "FULL_TIME",
        "description": "Seeking FastAPI Developer",
        "required_skills": ["Python", "FastAPI"],
        "status": "OPEN"
    }, headers=headers)
    assert r.status_code == 201
    job_id = r.json()["id"]
    
    # Trigger matching
    print("Triggering candidate matching run...")
    r = client.post(f"/api/v1/matching/jobs/{job_id}/run", headers=headers)
    assert r.status_code == 202
    
    # Check matching status (should be COMPLETED)
    r = client.get(f"/api/v1/matching/jobs/{job_id}/status", headers=headers)
    assert r.status_code == 200
    match_status = r.json()
    assert match_status["status"] == "COMPLETED"
    assert match_status["total_candidates_checked"] == 1
    assert match_status["total_matches_created"] == 1
    print("Success: Matching run completed.")
    
    # Verify match results are saved in DB
    mr = db.query(MatchResult).filter(MatchResult.job_id == job_id, MatchResult.candidate_id == cand_id).first()
    assert mr is not None
    assert mr.match_percentage == 92.0
    assert mr.vector_similarity_score is None # Must be None since Qdrant is disabled!
    assert mr.final_recommendation == "SHORTLIST"
    print(f"Success: MatchResult saved. vector_similarity_score is correctly None: {mr.vector_similarity_score}")
    
    # Verify recruitment events
    events = db.query(RecruitmentEvent).filter(RecruitmentEvent.job_id == job_id).all()
    event_types = [e.event_type for e in events]
    print(f"Matching Recorded Events: {event_types}")
    assert "MATCHING_STARTED" in event_types
    assert "CANDIDATE_MATCH_CREATED" in event_types
    assert "MATCHING_COMPLETED" in event_types
    assert "JOB_EMBEDDING_CREATED" not in event_types # Skipped since Qdrant is disabled
    print("Success: Match event milestones correctly recorded.")

    # 4. Test Failed Gemini Parsing Response sets status to FAILED
    print("\n4. Testing Failed Gemini Response parsing status FAILED...")
    # Trigger parse failure by making Gemini mock throw an error
    mock_parse_gemini.side_effect = Exception("Gemini API Overloaded!")
    
    # Trigger presigned url for a second candidate
    r = client.post("/api/v1/uploads/presigned-url", json={
        "first_name": "Failed",
        "last_name": "Candidate",
        "email": "failed.cand@gmail.com",
        "file_name": "resume.docx",
        "file_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    }, headers=headers)
    assert r.status_code == 200
    failed_cand_id = r.json()["candidate_id"]
    failed_resume_id = r.json()["resume_id"]
    
    # Write mock docx file
    failed_mock_file = f"uploads/resumes/{failed_resume_id}/resume.docx"
    os.makedirs(os.path.dirname(failed_mock_file), exist_ok=True)
    create_test_docx(failed_mock_file)
        
    # Trigger parse (should run synchronously and fail due to our mock side-effect)
    print("Triggering failing parse...")
    r = client.post(f"/api/v1/resumes/{failed_resume_id}/parse", headers=headers)
    assert r.status_code == 202
    
    # Check resume status (should be FAILED)
    r = client.get(f"/api/v1/resumes/{failed_resume_id}/status", headers=headers)
    assert r.status_code == 200
    failed_status = r.json()
    assert failed_status["status"] == "FAILED"
    assert "Gemini API Overloaded!" in failed_status["failure_reason"]
    print(f"Success: Failed parse set status to FAILED. Reason: '{failed_status['failure_reason']}'")
    
    # Check that events log RESUME_PARSING_FAILED
    failed_events = db.query(RecruitmentEvent).filter(RecruitmentEvent.resume_id == failed_resume_id).all()
    failed_event_types = [e.event_type for e in failed_events]
    print(f"Failed candidate events: {failed_event_types}")
    assert "RESUME_PARSING_FAILED" in failed_event_types
    print("Success: RESUME_PARSING_FAILED logged.")

    db.close()
    
    # Clean up test SQLite database and mock files
    try:
        os.remove("verify_mvp_no_qdrant.db")
        # Clean uploads files created
        shutil.rmtree("uploads/resumes")
        print("\n5. Test database and uploads folders cleaned up.")
    except Exception as e:
        print(f"Cleanup warning: {str(e)}")

    print("\n--- ALL REVISED MVP (NO QDRANT) INTEGRATION TESTS PASSED SUCCESSFULLY! ---")

if __name__ == "__main__":
    run_tests()
